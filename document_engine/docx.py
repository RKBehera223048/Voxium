import os
import re
from typing import Optional
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from .exporter import BaseExporter

class DocxExporter(BaseExporter):
    """Exports markdown content to DOCX."""
    
    def export(self, content: str, output_path: str, title: Optional[str] = None) -> str:
        doc = docx.Document()
        
        if title:
            heading = doc.add_heading(title, 0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        lines = content.split('\n')
        in_code_block = False
        code_lines = []
        
        for line in lines:
            if line.startswith('```'):
                if in_code_block:
                    p = doc.add_paragraph('\n'.join(code_lines))
                    p.style = 'Normal'
                    in_code_block = False
                    code_lines = []
                else:
                    in_code_block = True
                continue
                
            if in_code_block:
                code_lines.append(line)
                continue
                
            heading_match = re.match(r'^(#{1,6})\s+(.*)$', line)
            if heading_match:
                level = len(heading_match.group(1))
                doc.add_heading(heading_match.group(2), level=level)
                continue
                
            if line.startswith('- ') or line.startswith('* '):
                doc.add_paragraph(line[2:], style='List Bullet')
                continue
                
            list_match = re.match(r'^\d+\.\s+(.*)$', line)
            if list_match:
                doc.add_paragraph(list_match.group(1), style='List Number')
                continue
                
            if line.strip() == '':
                continue
                
            p = doc.add_paragraph()
            text = line
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text) 
            text = re.sub(r'\*(.*?)\*', r'\1', text)
            p.add_run(text)
            
        doc.save(output_path)
        return os.path.abspath(output_path)
